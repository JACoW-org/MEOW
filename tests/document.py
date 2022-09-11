from anyio import run


async def docx_main():
    from docx import Document

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_page_break()

    document.save("demo.docx")

    print('document.save("demo.docx")')


async def odt_main():
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties
    from odf.text import H, P, Span

    text_doc = OpenDocumentText()
    # Styles
    s = text_doc.styles
    h1style = Style(name="Heading 1", family="paragraph")
    h1style.addElement(TextProperties(attributes={'fontsize': "24pt", 'fontweight': "bold"}))
    s.addElement(h1style)
    # An automatic style
    bold_style = Style(name="Bold", family="text")
    bold_prop = TextProperties(fontweight="bold")
    bold_style.addElement(bold_prop)
    text_doc.automaticstyles.addElement(bold_style)
    # Text
    h = H(outlinelevel=1, stylename=h1style, text="My first text")
    text_doc.text.addElement(h)
    p = P(text="Hello world. ")
    bold_part = Span(stylename=bold_style, text="This part is bold. ")
    p.addElement(bold_part)
    p.addText("This is after bold.")
    text_doc.text.addElement(p)
    # text_doc.save("myfirstdocument.odt")

    print(text_doc.xml())

    with open('test/test.xml.odt', 'w') as fd:
        fd.write(str(text_doc.xml(), 'utf-8'))

    print('text_doc.save("myfirstdocument.odt")')


async def rtf_main():
    from PyRTF.Renderer import Renderer
    from PyRTF.Elements import Document
    from PyRTF.document.section import Section

    doc = Document()

    section = Section()
    doc.Sections.append(section)

    #    text can be added directly to the section
    #    a paragraph object is create as needed
    section.append('Example 1')

    #    blank paragraphs are just empty strings
    section.append('')

    #    a lot of useful documents can be created
    #    with little more than this
    section.append('A lot of useful documents can be created '
                   'in this way, more advance formating is available '
                   'but a lot of users just want to see their data come out '
                   'in something other than a text file.')

    import io

    # open('test.rtf', 'w')
    with io.StringIO() as f:
        Renderer().Write(doc, f)


async def main():
    for x in range(1):
        await odt_main()
        # await docx_main()
        # await rtf_main()


run(main)
